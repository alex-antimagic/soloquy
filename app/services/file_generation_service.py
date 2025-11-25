"""
File Generation Service
Handles generation of PDF, CSV, and Excel files with Cloudinary storage
"""
import io
import csv
from datetime import datetime
from typing import Dict, List, Optional, Any
from flask import current_app

# PDF generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

# Excel generation
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Word document generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from app.services.cloudinary_service import upload_file


class FileGenerationService:
    """Service for generating various file types"""

    @staticmethod
    def generate_pdf(
        title: str,
        content_blocks: List[Dict[str, Any]],
        tenant_id: int,
        filename: Optional[str] = None,
        page_size: str = "letter"
    ) -> Dict[str, Any]:
        """
        Generate a PDF document and upload to Cloudinary.

        Args:
            title: Document title
            content_blocks: List of content blocks, each with:
                - type: 'heading', 'paragraph', 'table', 'spacer'
                - content: The content for that block
            tenant_id: Tenant ID for file organization
            filename: Optional custom filename (auto-generated if not provided)
            page_size: Page size ('letter' or 'a4')

        Returns:
            Dictionary with Cloudinary upload result and metadata

        Example content_blocks:
            [
                {'type': 'heading', 'content': 'Q4 Sales Report', 'level': 1},
                {'type': 'paragraph', 'content': 'This report covers Q4 2024 sales...'},
                {'type': 'spacer', 'height': 0.25},
                {'type': 'table', 'content': {
                    'headers': ['Product', 'Sales', 'Revenue'],
                    'data': [['Widget A', '100', '$10,000'], ...]
                }}
            ]
        """
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
                safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
                safe_title = safe_title.replace(' ', '_')[:50]
                filename = f"{safe_title}_{timestamp}.pdf"

            # Create PDF in memory
            buffer = io.BytesIO()

            # Set page size
            page_size_obj = letter if page_size == "letter" else A4
            doc = SimpleDocTemplate(buffer, pagesize=page_size_obj)

            # Container for PDF elements
            story = []

            # Get styles
            styles = getSampleStyleSheet()

            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=30,
                alignment=TA_CENTER
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                textColor=colors.HexColor('#333333'),
                spaceAfter=12
            )

            # Add title
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.3 * inch))

            # Process content blocks
            for block in content_blocks:
                block_type = block.get('type')
                content = block.get('content')

                if block_type == 'heading':
                    level = block.get('level', 2)
                    style = heading_style if level == 2 else styles['Heading3']
                    story.append(Paragraph(str(content), style))
                    story.append(Spacer(1, 0.1 * inch))

                elif block_type == 'paragraph':
                    story.append(Paragraph(str(content), styles['Normal']))
                    story.append(Spacer(1, 0.15 * inch))

                elif block_type == 'table':
                    headers = content.get('headers', [])
                    data = content.get('data', [])

                    # Build table data
                    table_data = [headers] if headers else []
                    table_data.extend(data)

                    # Create table
                    table = Table(table_data)

                    # Style table
                    table_style = TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4A5568')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 10),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F7FAFC')])
                    ])

                    table.setStyle(table_style)
                    story.append(table)
                    story.append(Spacer(1, 0.2 * inch))

                elif block_type == 'spacer':
                    height = block.get('height', 0.25)
                    story.append(Spacer(1, height * inch))

            # Build PDF
            doc.build(story)

            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.seek(0)

            # Upload to Cloudinary
            folder = f"tenant_{tenant_id}/generated"
            cloudinary_result = upload_file(buffer, folder=folder, filename=filename)

            # Return result with metadata
            return {
                'success': True,
                'cloudinary_url': cloudinary_result['secure_url'],
                'public_id': cloudinary_result['public_id'],
                'filename': filename,
                'file_size': cloudinary_result['bytes'],
                'file_type': 'pdf',
                'mime_type': 'application/pdf'
            }

        except Exception as e:
            current_app.logger.error(f"Failed to generate PDF: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }

    @staticmethod
    def generate_csv(
        headers: List[str],
        data: List[List[Any]],
        tenant_id: int,
        filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Generate a CSV file and upload to Cloudinary.

        Args:
            headers: List of column headers
            data: List of rows, each row is a list of values
            tenant_id: Tenant ID for file organization
            filename: Optional custom filename

        Returns:
            Dictionary with Cloudinary upload result and metadata
        """
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
                filename = f"export_{timestamp}.csv"

            # Create CSV in memory
            output = io.StringIO()
            writer = csv.writer(output)

            # Write headers
            if headers:
                writer.writerow(headers)

            # Write data
            for row in data:
                writer.writerow(row)

            # Convert to bytes
            csv_content = output.getvalue()
            csv_bytes = io.BytesIO(csv_content.encode('utf-8'))

            # Upload to Cloudinary
            folder = f"tenant_{tenant_id}/generated"
            cloudinary_result = upload_file(csv_bytes, folder=folder, filename=filename)

            # Return result with metadata
            return {
                'success': True,
                'cloudinary_url': cloudinary_result['secure_url'],
                'public_id': cloudinary_result['public_id'],
                'filename': filename,
                'file_size': cloudinary_result['bytes'],
                'file_type': 'csv',
                'mime_type': 'text/csv'
            }

        except Exception as e:
            current_app.logger.error(f"Failed to generate CSV: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }

    @staticmethod
    def generate_excel(
        sheets: List[Dict[str, Any]],
        tenant_id: int,
        filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Generate an Excel workbook and upload to Cloudinary.

        Args:
            sheets: List of sheet definitions, each with:
                - name: Sheet name
                - headers: List of column headers
                - data: List of rows (each row is a list of values)
                - column_widths: Optional dict mapping column letters to widths
            tenant_id: Tenant ID for file organization
            filename: Optional custom filename

        Returns:
            Dictionary with Cloudinary upload result and metadata

        Example sheets:
            [
                {
                    'name': 'Sales Data',
                    'headers': ['Product', 'Quantity', 'Revenue'],
                    'data': [['Widget A', 100, 10000], ...],
                    'column_widths': {'A': 20, 'B': 15, 'C': 15}
                }
            ]
        """
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
                filename = f"export_{timestamp}.xlsx"

            # Create workbook
            wb = Workbook()

            # Remove default sheet
            if 'Sheet' in wb.sheetnames:
                wb.remove(wb['Sheet'])

            # Define styles
            header_font = Font(bold=True, color="FFFFFF", size=12)
            header_fill = PatternFill(start_color="4A5568", end_color="4A5568", fill_type="solid")
            header_alignment = Alignment(horizontal="left", vertical="center")

            cell_alignment = Alignment(horizontal="left", vertical="center")
            border = Border(
                left=Side(style='thin', color='D3D3D3'),
                right=Side(style='thin', color='D3D3D3'),
                top=Side(style='thin', color='D3D3D3'),
                bottom=Side(style='thin', color='D3D3D3')
            )

            # Create sheets
            for sheet_config in sheets:
                sheet_name = sheet_config.get('name', 'Sheet1')
                headers = sheet_config.get('headers', [])
                data = sheet_config.get('data', [])
                column_widths = sheet_config.get('column_widths', {})

                # Create sheet
                ws = wb.create_sheet(title=sheet_name)

                # Write headers
                if headers:
                    for col_idx, header in enumerate(headers, start=1):
                        cell = ws.cell(row=1, column=col_idx, value=header)
                        cell.font = header_font
                        cell.fill = header_fill
                        cell.alignment = header_alignment
                        cell.border = border

                # Write data
                for row_idx, row_data in enumerate(data, start=2):
                    for col_idx, value in enumerate(row_data, start=1):
                        cell = ws.cell(row=row_idx, column=col_idx, value=value)
                        cell.alignment = cell_alignment
                        cell.border = border

                        # Alternate row colors
                        if row_idx % 2 == 0:
                            cell.fill = PatternFill(start_color="F7FAFC", end_color="F7FAFC", fill_type="solid")

                # Set column widths
                if column_widths:
                    for col_letter, width in column_widths.items():
                        ws.column_dimensions[col_letter].width = width
                else:
                    # Auto-size columns based on content
                    for col_idx in range(1, len(headers) + 1):
                        col_letter = get_column_letter(col_idx)
                        ws.column_dimensions[col_letter].width = 15

            # Save to bytes
            excel_bytes = io.BytesIO()
            wb.save(excel_bytes)
            excel_bytes.seek(0)

            # Upload to Cloudinary
            folder = f"tenant_{tenant_id}/generated"
            cloudinary_result = upload_file(excel_bytes, folder=folder, filename=filename)

            # Return result with metadata
            return {
                'success': True,
                'cloudinary_url': cloudinary_result['secure_url'],
                'public_id': cloudinary_result['public_id'],
                'filename': filename,
                'file_size': cloudinary_result['bytes'],
                'file_type': 'xlsx',
                'mime_type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            }

        except Exception as e:
            current_app.logger.error(f"Failed to generate Excel file: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }

    @staticmethod
    def generate_markdown(
        content: str,
        tenant_id: int,
        filename: Optional[str] = None,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Generate a Markdown file and upload to Cloudinary.

        Args:
            content: Markdown content (can include headers, lists, code blocks, etc.)
            tenant_id: Tenant ID for file organization
            filename: Optional custom filename
            title: Optional title to prepend as H1

        Returns:
            Dictionary with Cloudinary upload result and metadata

        Example:
            generate_markdown(
                content="## Introduction\\n\\nThis is a **bold** statement.",
                tenant_id=1,
                title="My Document"
            )
        """
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
                if title:
                    safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
                    safe_title = safe_title.replace(' ', '_')[:50]
                    filename = f"{safe_title}_{timestamp}.md"
                else:
                    filename = f"document_{timestamp}.md"

            # Build markdown content
            markdown_content = ""
            if title:
                markdown_content = f"# {title}\n\n"
            markdown_content += content

            # Convert to bytes
            md_bytes = io.BytesIO(markdown_content.encode('utf-8'))

            # Upload to Cloudinary
            folder = f"tenant_{tenant_id}/generated"
            cloudinary_result = upload_file(md_bytes, folder=folder, filename=filename)

            return {
                'success': True,
                'cloudinary_url': cloudinary_result['secure_url'],
                'public_id': cloudinary_result['public_id'],
                'filename': filename,
                'file_size': cloudinary_result['bytes'],
                'file_type': 'md',
                'mime_type': 'text/markdown'
            }

        except Exception as e:
            current_app.logger.error(f"Failed to generate Markdown file: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }

    @staticmethod
    def generate_docx(
        title: str,
        content_blocks: List[Dict[str, Any]],
        tenant_id: int,
        filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Generate a Word document (.docx) and upload to Cloudinary.

        Args:
            title: Document title
            content_blocks: List of content blocks, each with:
                - type: 'heading', 'paragraph', 'bullet_list', 'numbered_list', 'table'
                - content: The content for that block
                - level: For headings, 1-3
            tenant_id: Tenant ID for file organization
            filename: Optional custom filename

        Returns:
            Dictionary with Cloudinary upload result and metadata

        Example content_blocks:
            [
                {'type': 'heading', 'content': 'Introduction', 'level': 1},
                {'type': 'paragraph', 'content': 'This document covers...'},
                {'type': 'bullet_list', 'content': ['Item 1', 'Item 2', 'Item 3']},
                {'type': 'numbered_list', 'content': ['First step', 'Second step']},
                {'type': 'table', 'content': {
                    'headers': ['Name', 'Value'],
                    'data': [['Row 1', '100'], ['Row 2', '200']]
                }}
            ]
        """
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
                safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
                safe_title = safe_title.replace(' ', '_')[:50]
                filename = f"{safe_title}_{timestamp}.docx"

            # Create document
            doc = Document()

            # Add title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Process content blocks
            for block in content_blocks:
                block_type = block.get('type')
                content = block.get('content')

                if block_type == 'heading':
                    level = min(block.get('level', 1), 3)  # Max level 3
                    doc.add_heading(str(content), level=level)

                elif block_type == 'paragraph':
                    doc.add_paragraph(str(content))

                elif block_type == 'bullet_list':
                    if isinstance(content, list):
                        for item in content:
                            doc.add_paragraph(str(item), style='List Bullet')

                elif block_type == 'numbered_list':
                    if isinstance(content, list):
                        for item in content:
                            doc.add_paragraph(str(item), style='List Number')

                elif block_type == 'table':
                    headers = content.get('headers', [])
                    data = content.get('data', [])

                    # Create table with header row
                    num_cols = len(headers) if headers else (len(data[0]) if data else 0)
                    num_rows = (1 if headers else 0) + len(data)

                    if num_cols > 0 and num_rows > 0:
                        table = doc.add_table(rows=num_rows, cols=num_cols)
                        table.style = 'Table Grid'

                        # Add headers
                        if headers:
                            header_row = table.rows[0]
                            for idx, header in enumerate(headers):
                                cell = header_row.cells[idx]
                                cell.text = str(header)
                                # Make header bold
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True

                        # Add data
                        start_row = 1 if headers else 0
                        for row_idx, row_data in enumerate(data):
                            row = table.rows[start_row + row_idx]
                            for col_idx, value in enumerate(row_data):
                                if col_idx < num_cols:
                                    row.cells[col_idx].text = str(value)

                        # Add spacing after table
                        doc.add_paragraph()

            # Save to bytes
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)

            # Upload to Cloudinary
            folder = f"tenant_{tenant_id}/generated"
            cloudinary_result = upload_file(docx_bytes, folder=folder, filename=filename)

            return {
                'success': True,
                'cloudinary_url': cloudinary_result['secure_url'],
                'public_id': cloudinary_result['public_id'],
                'filename': filename,
                'file_size': cloudinary_result['bytes'],
                'file_type': 'docx',
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }

        except Exception as e:
            current_app.logger.error(f"Failed to generate Word document: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
